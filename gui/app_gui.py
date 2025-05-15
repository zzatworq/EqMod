import tkinter as tk
from tkinter import ttk, messagebox, filedialog
import tkinter.font as tkfont
import threading
import sys
import logging
import base64
import win32clipboard
from docx import Document
from docx.shared import Pt
import os
import json
from configs.settings import LOGGING_CONFIG
from utils.clipboard_utils import set_clipboard_html, get_clipboard_text, validate_base64
from utils.latex_utils import check_latex, find_latex_equations
from utils.image_utils import render_latex_to_image, image_to_bytes, is_image_empty
import time
import html

class LatexClipboardApp:
    def __init__(self, root):
        self.root = root
        self.root.title("LaTeX Clipboard Monitor")
        self.monitoring = False
        self.monitor_thread = None
        self.stop_event = threading.Event()
        self.last_images = []
        self.last_text = ""
        self.last_original_text = ""
        self.last_equations = None
        self.defaults_file = os.path.join("configs", "defaults.json")

        if not check_latex():
            messagebox.showerror("LaTeX Not Found",
                                 "A LaTeX distribution (e.g., MiKTeX) with latex and dvipng is required. Please install MiKTeX and try again.")
            sys.exit(1)

        self.load_defaults()
        self.root.state('normal')
        self.root.attributes('-topmost', True)
        self.root.update()
        self.root.attributes('-topmost', False)
        self.root.focus_force()

        self.create_gui()
        logging.info("Application initialized. Using white text on transparent background.")

    def load_defaults(self):
        """Load default settings from defaults.json or set fallback defaults."""
        default_settings = {
            "mode": "Matplotlib",
            "text_color": "white",
            "font_size": "12",
            "dpi": "300",
            "only_images": False
        }
        try:
            if os.path.exists(self.defaults_file):
                with open(self.defaults_file, 'r') as f:
                    defaults = json.load(f)
                # Validate loaded defaults
                default_settings.update({
                    k: v for k, v in defaults.items()
                    if k in default_settings and isinstance(v, type(default_settings[k]))
                })
                logging.info(f"Loaded defaults from {self.defaults_file}: {default_settings}")
            else:
                logging.info(f"No defaults file found at {self.defaults_file}, using fallback defaults")
        except Exception as e:
            logging.error(f"Failed to load defaults: {e}, using fallback defaults")
        self.default_settings = default_settings

    def save_defaults(self, settings):
        """Save default settings to defaults.json."""
        try:
            os.makedirs(os.path.dirname(self.defaults_file), exist_ok=True)
            with open(self.defaults_file, 'w') as f:
                json.dump(settings, f, indent=4)
            logging.info(f"Saved defaults to {self.defaults_file}: {settings}")
        except Exception as e:
            logging.error(f"Failed to save defaults: {e}")
            messagebox.showerror("Save Defaults Failed", f"Error saving defaults: {e}")

    def create_gui(self):
        style = ttk.Style()
        style.configure("TLabel", font=("Arial", 12))
        style.configure("TButton", font=("Arial", 12))
        style.configure("TCheckbutton", font=("Arial", 12))
        style.configure("TSpinbox", font=("Arial", 12))
        style.configure("TMenubutton", font=("Arial", 12))
        style.configure("Note.TLabel", font=("Arial", 12, "bold"))
        style.configure("Section.TLabel", font=("Arial", 14, "bold"))

        menu_font = tkfont.Font(family="Arial", size=12)

        # Main container
        main_frame = ttk.Frame(self.root, padding="10")
        main_frame.grid(row=0, column=0, sticky="nsew")
        self.root.columnconfigure(0, weight=1)
        self.root.rowconfigure(0, weight=1)

        # Settings Section
        settings_frame = ttk.LabelFrame(main_frame, text="Configuration", padding="5")
        settings_frame.grid(row=0, column=0, sticky="ew", pady=5)
        settings_frame.columnconfigure(1, weight=1)

        # Mode selector
        ttk.Label(settings_frame, text="Render Mode:").grid(row=0, column=0, padx=5, pady=5, sticky="e")
        self.mode_var = tk.StringVar(value=self.default_settings["mode"])
        modes = ["Matplotlib", "Standalone"]
        self.mode_menu = ttk.OptionMenu(settings_frame, self.mode_var, self.default_settings["mode"], *modes)
        self.mode_menu["menu"].configure(font=menu_font)
        self.mode_menu.grid(row=0, column=1, padx=5, pady=5, sticky="w")
        self.add_tooltip(self.mode_menu, "Select rendering mode: Matplotlib (Python-based) or Standalone (uses latex/dvipng)")

        # Text color
        ttk.Label(settings_frame, text="Text Color:").grid(row=1, column=0, padx=5, pady=5, sticky="e")
        self.color_var = tk.StringVar(value=self.default_settings["text_color"])
        colors = ["white", "black", "red", "blue", "green"]
        self.color_menu = ttk.OptionMenu(settings_frame, self.color_var, self.default_settings["text_color"], *colors)
        self.color_menu["menu"].configure(font=menu_font)
        self.color_menu.grid(row=1, column=1, padx=5, pady=5, sticky="w")
        self.add_tooltip(self.color_menu, "Set text color for rendered equations and surrounding text")

        # Font size
        ttk.Label(settings_frame, text="Font Size:").grid(row=2, column=0, padx=5, pady=5, sticky="e")
        self.font_size_var = tk.StringVar(value=self.default_settings["font_size"])
        self.font_size_spin = ttk.Spinbox(settings_frame, from_=10, to=50, width=10, textvariable=self.font_size_var, font=("Arial", 12))
        self.font_size_spin.grid(row=2, column=1, padx=5, pady=5, sticky="w")
        self.add_tooltip(self.font_size_spin, "Set font size (10-50 pt) for text in output")

        # DPI
        ttk.Label(settings_frame, text="DPI:").grid(row=3, column=0, padx=5, pady=5, sticky="e")
        self.dpi_var = tk.StringVar(value=self.default_settings["dpi"])
        self.dpi_spin = ttk.Spinbox(settings_frame, from_=100, to=600, width=10, textvariable=self.dpi_var, font=("Arial", 12))
        self.dpi_spin.grid(row=3, column=1, padx=5, pady=5, sticky="w")
        self.add_tooltip(self.dpi_spin, "Set image resolution (100-600 DPI) for rendered equations")

        # Only images checkbox
        self.only_images_var = tk.BooleanVar(value=self.default_settings["only_images"])
        self.only_images_check = ttk.Checkbutton(settings_frame, text="Only Images (Ignore Text)", variable=self.only_images_var)
        self.only_images_check.grid(row=4, column=0, columnspan=2, padx=5, pady=5, sticky="w")
        self.add_tooltip(self.only_images_check, "Output only equation images, ignoring surrounding text")

        # Actions Section
        actions_frame = ttk.LabelFrame(main_frame, text="Actions", padding="5")
        actions_frame.grid(row=1, column=0, sticky="ew", pady=5)

        self.toggle_button = ttk.Button(actions_frame, text="Start Monitoring", command=self.toggle_monitoring)
        self.toggle_button.grid(row=0, column=0, padx=5, pady=5)

        self.test_button = ttk.Button(actions_frame, text="Test Render", command=self.test_render)
        self.test_button.grid(row=0, column=1, padx=5, pady=5)

        self.save_button = ttk.Button(actions_frame, text="Save as DOCX", command=self.save_as_docx)
        self.save_button.grid(row=0, column=2, padx=5, pady=5)

        self.defaults_button = ttk.Button(actions_frame, text="Defaults", command=self.open_defaults_dialog)
        self.defaults_button.grid(row=0, column=3, padx=5, pady=5)

        # Input/Output Section
        io_frame = ttk.LabelFrame(main_frame, text="Input & Output", padding="5")
        io_frame.grid(row=2, column=0, sticky="nsew", pady=5)
        io_frame.columnconfigure(0, weight=1)
        io_frame.rowconfigure(1, weight=1)

        ttk.Label(io_frame, text="Input Text:").grid(row=0, column=0, padx=5, pady=5, sticky="nw")
        self.text_input = tk.Text(io_frame, height=5, width=50, font=("Arial", 12))
        self.text_input.grid(row=1, column=0, padx=5, pady=5, sticky="nsew")
        self.add_tooltip(self.text_input, "Enter text with LaTeX equations (e.g., \\[E=mc^2\\]) to render")

        self.render_button = ttk.Button(io_frame, text="Render Input", command=self.render_input_text)
        self.render_button.grid(row=2, column=0, padx=5, pady=5, sticky="ew")

        # Status and Notes
        ttk.Label(io_frame, text="Note: White text may be invisible on white backgrounds in some applications.", foreground="red", style="Note.TLabel").grid(row=3, column=0, padx=5, pady=5, sticky="w")
        self.status_var = tk.StringVar(value="Stopped")
        ttk.Label(io_frame, textvariable=self.status_var, foreground="red").grid(row=4, column=0, padx=5, pady=5, sticky="w")

        self.root.protocol("WM_DELETE_WINDOW", self.on_closing)

    def open_defaults_dialog(self):
        """Open a dialog to configure default settings."""
        dialog = tk.Toplevel(self.root)
        dialog.title("Set Default Settings")
        dialog.transient(self.root)
        dialog.grab_set()
        dialog.resizable(False, False)

        frame = ttk.Frame(dialog, padding="15")
        frame.grid(row=0, column=0, sticky="nsew")
        dialog.columnconfigure(0, weight=1)
        dialog.rowconfigure(0, weight=1)

        # Mode
        ttk.Label(frame, text="Default Render Mode:").grid(row=0, column=0, padx=10, pady=10, sticky="e")
        mode_var = tk.StringVar(value=self.default_settings["mode"])
        modes = ["Matplotlib", "Standalone"]
        ttk.OptionMenu(frame, mode_var, self.default_settings["mode"], *modes).grid(row=0, column=1, padx=10, pady=10, sticky="w")

        # Text color
        ttk.Label(frame, text="Default Text Color:").grid(row=1, column=0, padx=10, pady=10, sticky="e")
        color_var = tk.StringVar(value=self.default_settings["text_color"])
        colors = ["white", "black", "red", "blue", "green"]
        ttk.OptionMenu(frame, color_var, self.default_settings["text_color"], *colors).grid(row=1, column=1, padx=10, pady=10, sticky="w")

        # Font size
        ttk.Label(frame, text="Default Font Size:").grid(row=2, column=0, padx=10, pady=10, sticky="e")
        font_size_var = tk.StringVar(value=self.default_settings["font_size"])
        ttk.Spinbox(frame, from_=10, to=50, width=10, textvariable=font_size_var).grid(row=2, column=1, padx=10, pady=10, sticky="w")

        # DPI
        ttk.Label(frame, text="Default DPI:").grid(row=3, column=0, padx=10, pady=10, sticky="e")
        dpi_var = tk.StringVar(value=self.default_settings["dpi"])
        ttk.Spinbox(frame, from_=100, to=600, width=10, textvariable=dpi_var).grid(row=3, column=1, padx=10, pady=10, sticky="w")

        # Only images
        only_images_var = tk.BooleanVar(value=self.default_settings["only_images"])
        ttk.Checkbutton(frame, text="Default Only Images", variable=only_images_var).grid(row=4, column=0, columnspan=2, padx=10, pady=10, sticky="w")

        # Buttons
        def save():
            try:
                font_size = int(font_size_var.get())
                dpi = int(dpi_var.get())
                if not (10 <= font_size <= 50):
                    raise ValueError("Font size must be between 10 and 50")
                if not (100 <= dpi <= 600):
                    raise ValueError("DPI must be between 100 and 600")
                new_defaults = {
                    "mode": mode_var.get(),
                    "text_color": color_var.get(),
                    "font_size": str(font_size),
                    "dpi": str(dpi),
                    "only_images": only_images_var.get()
                }
                self.default_settings = new_defaults
                self.save_defaults(new_defaults)
                # Update main GUI
                self.mode_var.set(new_defaults["mode"])
                self.color_var.set(new_defaults["text_color"])
                self.font_size_var.set(new_defaults["font_size"])
                self.dpi_var.set(new_defaults["dpi"])
                self.only_images_var.set(new_defaults["only_images"])
                messagebox.showinfo("Defaults Saved", "Default settings updated successfully.")
                dialog.destroy()
            except ValueError as e:
                messagebox.showerror("Invalid Input", str(e))

        def cancel():
            dialog.destroy()

        ttk.Button(frame, text="Save", command=save).grid(row=5, column=0, padx=10, pady=15)
        ttk.Button(frame, text="Cancel", command=cancel).grid(row=5, column=1, padx=10, pady=15)

        # Auto-size dialog
        dialog.update_idletasks()
        width = frame.winfo_reqwidth() + 20  # Add padding
        height = frame.winfo_reqheight() + 20
        x = self.root.winfo_rootx() + (self.root.winfo_width() - width) // 2
        y = self.root.winfo_rooty() + (self.root.winfo_height() - height) // 2
        dialog.geometry(f"{width}x{height}+{x}+{y}")

    def add_tooltip(self, widget, text):
        """Add a tooltip to a widget."""
        tooltip = tk.Toplevel(self.root)
        tooltip.wm_overrideredirect(True)
        tooltip.wm_geometry("+0+0")
        label = ttk.Label(tooltip, text=text, background="lightyellow", relief="solid", borderwidth=1, font=("Arial", 10))
        label.pack()

        def show(event):
            x = widget.winfo_rootx() + 20
            y = widget.winfo_rooty() + widget.winfo_height()
            tooltip.wm_geometry(f"+{x}+{y}")
            tooltip.deiconify()

        def hide(event):
            tooltip.withdraw()

        widget.bind("<Enter>", show)
        widget.bind("<Leave>", hide)
        tooltip.withdraw()

    def validate_inputs(self):
        try:
            font_size = int(self.font_size_var.get())
            if not 10 <= font_size <= 50:
                raise ValueError("Font size must be between 10 and 50")
            dpi = int(self.dpi_var.get())
            if not 100 <= dpi <= 600:
                raise ValueError("DPI must be between 100 and 600")
            return True
        except ValueError as e:
            logging.error(f"Input validation failed: {e}")
            messagebox.showerror("Invalid Input", str(e))
            return False

    def toggle_monitoring(self):
        if not self.monitoring:
            if not self.validate_inputs():
                return
            self.monitoring = True
            self.status_var.set("Monitoring")
            self.toggle_button.configure(text="Stop Monitoring")
            self.mode_menu.configure(state="disabled")
            self.color_menu.configure(state="disabled")
            self.font_size_spin.configure(state="disabled")
            self.dpi_spin.configure(state="disabled")
            self.test_button.configure(state="disabled")
            self.save_button.configure(state="disabled")
            self.defaults_button.configure(state="disabled")
            self.render_button.configure(state="disabled")
            self.text_input.configure(state="disabled")
            self.stop_event.clear()
            self.root.after(1000, self.start_monitor_thread)
            logging.info("Started clipboard monitoring")
        else:
            self.monitoring = False
            self.status_var.set("Stopped")
            self.toggle_button.configure(text="Start Monitoring")
            self.mode_menu.configure(state="normal")
            self.color_menu.configure(state="normal")
            self.font_size_spin.configure(state="normal")
            self.dpi_spin.configure(state="normal")
            self.test_button.configure(state="normal")
            self.save_button.configure(state="normal")
            self.defaults_button.configure(state="normal")
            self.render_button.configure(state="normal")
            self.text_input.configure(state="normal")
            self.stop_event.set()
            logging.info("Stopped clipboard monitoring")

    def start_monitor_thread(self):
        self.monitor_thread = threading.Thread(target=self.monitor_clipboard, daemon=True)
        self.monitor_thread.start()

    def render_input_text(self):
        if not self.validate_inputs():
            return
        text = self.text_input.get("1.0", tk.END).strip()
        if not text:
            messagebox.showwarning("No Input", "Please enter text to render.")
            return
        logging.info(f"Rendering input text: {text[:100]}...")
        try:
            equations = find_latex_equations(text)
            images = []
            for i, eq in enumerate(equations['equations'], 1):
                logging.info(f"Rendering equation {i}: {eq[:50]}...")
                img = render_latex_to_image(eq, self.color_var.get(), int(self.font_size_var.get()), int(self.dpi_var.get()), mode=self.mode_var.get())
                if img:
                    images.append(img)
                    debug_path = f"input_eq_{len(images)-1}.png"
                    img.save(debug_path, format='PNG')
                    logging.info(f"Saved input debug image: {debug_path}")
                else:
                    logging.warning(f"Failed to render equation {i}: {eq[:50]}...")
            if images:
                self.copy_images(images, original_text=text, equations=equations)
                self.status_var.set(f"Copied {len(images)} images with text")
                messagebox.showinfo("Render Input", f"Copied {len(images)} images with text as HTML")
            else:
                self.status_var.set("No valid images rendered")
                messagebox.showerror("Render Input", "Failed to render any valid images")
        except Exception as e:
            logging.error(f"Input render failed: {e}")
            messagebox.showerror("Render Input Failed", f"Error rendering input equations: {e}")
            self.status_var.set("Render input failed")

    def test_render(self):
        if not self.validate_inputs():
            return
        from templates.test_string import TEST_STRING
        test_text = TEST_STRING
        logging.info(f"Rendering test text: {test_text[:100]}...")
        try:
            equations = find_latex_equations(test_text)
            images = []
            for i, eq in enumerate(equations['equations'], 1):
                logging.info(f"Rendering equation {i}: {eq[:50]}...")
                img = render_latex_to_image(eq, self.color_var.get(), int(self.font_size_var.get()), int(self.dpi_var.get()), mode=self.mode_var.get())
                if img:
                    images.append(img)
                    debug_path = f"./cache-and-logs/test_eq_{len(images)-1}.png"
                    img.save(debug_path, format='PNG')
                    logging.info(f"Saved test debug image: {debug_path}")
                else:
                    logging.warning(f"Failed to render equation {i}: {eq[:50]}...")
            if images:
                self.copy_images(images, test_mode=True, original_text=test_text, equations=equations)
                self.status_var.set(f"Copied {len(images)} images with text")
                messagebox.showinfo("Test Render", f"Copied {len(images)} images with text as HTML")
            else:
                self.status_var.set("No valid test images rendered")
                messagebox.showerror("Test Render", "Failed to render any valid images")
        except Exception as e:
            logging.error(f"Test render failed: {e}")
            messagebox.showerror("Test Render Failed", f"Error rendering test equations: {e}")
            self.status_var.set("Test render failed")

    def save_as_docx(self):
        if not self.last_images:
            messagebox.showwarning("No Images", "No images available to save. Run Test Render, Render Input, or monitor clipboard first.")
            return
        file_path = filedialog.asksaveasfilename(
            title="Save as DOCX",
            defaultextension=".docx",
            filetypes=[("Word Documents", "*.docx"), ("All Files", "*.*")]
        )
        if not file_path:
            return
        try:
            doc = Document()
            only_images = self.only_images_var.get()
            has_text = bool(self.last_original_text)
            has_equations = bool(self.last_equations and self.last_equations.get('matches'))
            logging.info(f"Saving DOCX: only_images={only_images}, text_length={len(self.last_original_text) if has_text else 0}, equations={len(self.last_equations['matches']) if has_equations else 0}, images={len(self.last_images)}")
            if has_equations:
                for i, match in enumerate(self.last_equations['matches'], 1):
                    logging.info(f"Equation match {i}: {match['equation'][:50]}... (raw: {match['raw_match'][:50]}...)")
            if only_images or not has_text or not has_equations:
                reasons = []
                if only_images:
                    reasons.append("'Only Images' is checked")
                if not has_text:
                    reasons.append("no text available")
                if not has_equations:
                    reasons.append("no equations available")
                logging.warning(f"Saving only images due to: {', '.join(reasons)}")
                messagebox.showinfo("Saving Images Only",
                                    f"Saving only equation images because: {', '.join(reasons)}.")
                for i, img in enumerate(self.last_images):
                    img_path = f"temp_eq_{i}.png"
                    img.save(img_path, format='PNG')
                    doc.add_picture(img_path, width=Pt(300))
                    os.remove(img_path)
            else:
                logging.info("Saving text and images")
                last_pos = 0
                img_index = 0
                font_size = int(self.font_size_var.get())
                for i, match in enumerate(self.last_equations['matches'], 1):
                    start, end = match['start'], match['end']
                    text_segment = self.last_original_text[last_pos:start]
                    if text_segment.strip():
                        logging.info(f"Adding text segment: {text_segment[:50]}...")
                        p = doc.add_paragraph()
                        p.add_run(text_segment).font.size = Pt(font_size)
                    if img_index < len(self.last_images):
                        img = self.last_images[img_index]
                        img_path = f"temp_eq_{img_index}.png"
                        img.save(img_path, format='PNG')
                        logging.info(f"Adding image {img_index + 1}/{len(self.last_images)} for equation {i}: {match['equation'][:50]}...")
                        p = doc.add_paragraph()
                        r = p.add_run()
                        r.add_picture(img_path, width=Pt(300))
                        os.remove(img_path)
                        img_index += 1
                    else:
                        logging.warning(f"No image available for equation {i}: {match['equation'][:50]}... (raw: {match['raw_match'][:50]}...)")
                    last_pos = end
                remaining_text = self.last_original_text[last_pos:]
                text_segment = remaining_text.strip()
                if text_segment:
                    logging.info(f"Adding final text segment: {text_segment[:50]}...")
                    p = doc.add_paragraph()
                    p.add_run(text_segment).font.size = Pt(font_size)
            doc.save(file_path)
            logging.info(f"Saved DOCX to {file_path}")
            messagebox.showinfo("Save Successful", f"Saved content to {file_path}")
            try:
                time.sleep(0.5)
                os.startfile(file_path)
                logging.info(f"Opened DOCX file: {file_path}")
            except Exception as e:
                logging.error(f"Failed to open DOCX file {file_path}: {e}")
                messagebox.showwarning("Open Failed", f"Saved file, but failed to open {file_path}: {e}")
        except Exception as e:
            logging.error(f"Failed to save DOCX: {e}")
            messagebox.showerror("Save Failed", f"Error saving DOCX file: {e}")

    def on_closing(self):
        if self.monitoring:
            self.stop_event.set()
            if self.monitor_thread:
                self.monitor_thread.join(timeout=1.0)
        self.root.destroy()
        logging.info("Application closed")

    def copy_images(self, images, test_mode=False, original_text="", equations=None):
        if not images:
            logging.info("No images to copy")
            return
        self.last_images = []
        self.last_original_text = original_text
        self.last_equations = equations or {'equations': [], 'matches': []}
        logging.info(f"Copying images: test_mode={test_mode}, text_length={len(original_text)}, equations={len(equations['matches']) if equations else 0}, images={len(images)}")

        html_content = ""
        text_color = self.color_var.get()
        font_size = int(self.font_size_var.get())

        # Validate text_color to ensure it's a valid CSS color
        valid_colors = ["white", "black", "red", "blue", "green", "yellow", "purple", "orange"]
        if text_color not in valid_colors:
            logging.warning(f"Invalid text color '{text_color}', defaulting to 'black'")
            text_color = "black"

        # Construct CSS style block as a single string
        style = (
            "<style>"
            "body {"
            f"  color: {text_color};"
            "  font-family: Arial, sans-serif;"
            f"  font-size: {font_size}pt;"
            "  line-height: 1.5;"
            "}"
            "p, div, span {"
            f"  color: {text_color} !important;"
            "}"
            "img {"
            "  vertical-align: middle;"
            "  margin: 2px 0;"
            "}"
            "</style>"
        )

        html_content += style
        logging.debug(f"CSS style block: {style}")

        if test_mode or original_text:
            if not equations or not equations.get('matches'):
                logging.warning("No valid equations provided for text replacement")
                for i, img in enumerate(images):
                    if img is None or is_image_empty(img):
                        logging.warning(f"Skipping empty/null image at index {i}")
                        continue
                    try:
                        bytes_data = image_to_bytes(img)
                        b64_data = base64.b64encode(bytes_data).decode('ascii')
                        if validate_base64(b64_data):
                            self.last_images.append(img)
                            html_content += f'<img src="data:image/png;base64,{b64_data}" style="vertical-align: middle; margin: 2px 0;">'
                            debug_path = f"./cache-and-logs/debug_eq_{i}.png"
                            img.save(debug_path, format='PNG')
                            logging.info(f"Saved debug image: {debug_path}")
                        else:
                            logging.warning(f"Invalid base64 data for image {i}")
                    except Exception as e:
                        logging.error(f"Error processing image {i}: {e}")
                if html_content and self.last_images:
                    try:
                        logging.debug(f"Setting clipboard HTML: {html_content[:200]}...")
                        set_clipboard_html(html_content)
                        self.status_var.set(f"Copied {len(self.last_images)} images")
                        logging.info(f"Copied {len(self.last_images)} images to clipboard")
                    except Exception as e:
                        logging.error(f"Failed to set clipboard HTML: {e}")
                        self.status_var.set("Error copying images")
                return
            valid_images = []
            valid_b64_data = []
            for i, img in enumerate(images):
                if img is None or is_image_empty(img):
                    logging.warning(f"Skipping empty/null image at index {i}")
                    continue
                try:
                    if img.width == 0 or img.height == 0:
                        logging.warning(f"Skipping zero-dimension image at index {i}")
                        continue
                    bytes_data = image_to_bytes(img)
                    b64_data = base64.b64encode(bytes_data).decode('ascii')
                    if not validate_base64(b64_data):
                        logging.warning(f"Invalid base64 data for image {i}")
                        continue
                    valid_images.append(img)
                    valid_b64_data.append(b64_data)
                    self.last_images.append(img)
                    debug_path = f"./cache-and-logs/debug_eq_{i}.png"
                    img.save(debug_path, format='PNG')
                    logging.info(f"Saved debug image: {debug_path}")
                except Exception as e:
                    logging.error(f"Error processing image {i}: {e}")
                    continue
            if not valid_images:
                logging.warning("No valid images to copy to clipboard")
                self.status_var.set("No valid images")
                return
            if self.only_images_var.get():
                for b64_data in valid_b64_data:
                    html_content += f'<img src="data:image/png;base64,{b64_data}" style="vertical-align: middle; margin: 2px 0;">'
            else:
                last_pos = 0
                img_index = 0
                for match in equations['matches']:
                    start, end = match['start'], match['end']
                    text_segment = original_text[last_pos:start]
                    text_segment = html.escape(text_segment).replace('\r\n', '<br>').replace('\n', '<br>')
                    html_content += f'<span>{text_segment}</span>'
                    if img_index < len(valid_b64_data):
                        html_content += f'<img src="data:image/png;base64,{valid_b64_data[img_index]}" style="vertical-align: middle; margin: 2px 0;">'
                        img_index += 1
                    last_pos = end
                text_segment = html.escape(original_text[last_pos:]).replace('\r\n', '<br>').replace('\n', '<br>')
                html_content += f'<span>{text_segment}</span>'
        else:
            for i, img in enumerate(images):
                if img is None or is_image_empty(img):
                    logging.warning(f"Skipping empty/null image at index {i}")
                    continue
                try:
                    if img.width == 0 or img.height == 0:
                        logging.warning(f"Skipping zero-dimension image at index {i}")
                        continue
                    bytes_data = image_to_bytes(img)
                    b64_data = base64.b64encode(bytes_data).decode('ascii')
                    if not validate_base64(b64_data):
                        logging.warning(f"Invalid base64 data for image {i}")
                        continue
                    html_content += f'<img src="data:image/png;base64,{b64_data}" style="vertical-align: middle; margin: 2px 0;">'
                    if i < len(images) - 1:
                        html_content += '<br>'
                    self.last_images.append(img)
                    debug_path = f"./cache-and-logs/debug_eq_{i}.png"
                    img.save(debug_path, format='PNG')
                    logging.info(f"Saved debug image: {debug_path}")
                except Exception as e:
                    logging.error(f"Error processing image {i}: {e}")
                    continue
        if html_content and self.last_images:
            try:
                logging.debug(f"Setting clipboard HTML: {html_content[:200]}...")
                set_clipboard_html(html_content)
                self.status_var.set(f"Copied {len(self.last_images)} images with text")
                logging.info(f"Successfully copied {len(self.last_images)} images with text to clipboard")
            except Exception as e:
                logging.error(f"Failed to set clipboard HTML: {e}")
                self.status_var.set("Error copying images")
        else:
            logging.warning("No valid images to copy to clipboard")
            self.status_var.set("No valid images")

    def monitor_clipboard(self):
        last_sequence = None
        processed_text = None
        while not self.stop_event.is_set():
            try:
                current_sequence = win32clipboard.GetClipboardSequenceNumber()
                if current_sequence != last_sequence:
                    last_sequence = current_sequence
                    text = get_clipboard_text()
                    if text and text != processed_text:
                        logging.info(f"New clipboard content detected: {text[:100]}...")
                        equations = find_latex_equations(text)
                        logging.info(f"Found {len(equations['equations'])} equations")
                        if equations['equations']:
                            images = []
                            for i, eq in enumerate(equations['equations'], 1):
                                logging.info(f"Processing equation {i}: {eq[:50]}...")
                                try:
                                    img = render_latex_to_image(eq, self.color_var.get(), int(self.font_size_var.get()), int(self.dpi_var.get()), mode=self.mode_var.get())
                                    if img:
                                        images.append(img)
                                    else:
                                        logging.warning(f"Equation {i} produced no valid image: {eq[:50]}...")
                                except Exception as e:
                                    logging.error(f"Skipping equation {i} due to rendering error: {e}")
                                    continue
                            if images:
                                self.copy_images(images, original_text=text, equations=equations)
                            else:
                                self.status_var.set("No valid images rendered")
                        else:
                            logging.info("No valid LaTeX equations found in clipboard content")
                            self.status_var.set("No equations found")
                    elif not text:
                        logging.info("Clipboard is empty or contains non-text data")
                        self.status_var.set("Clipboard empty")
                time.sleep(1)
            except Exception as e:
                logging.error(f"Error in clipboard monitoring loop: {e}")
                self.status_var.set("Error in monitoring")
                time.sleep(1)